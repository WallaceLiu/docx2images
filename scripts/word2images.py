#!/usr/bin/env python3
"""
Word → 多张图片 (手机端微头条专用)
====================================
用法:
    python3 word2images.py <input.docx>
    python3 word2images.py <input.docx> --output-dir ./output --font Chinese

依赖:
    pip3 install python-docx Pillow

手机端优化:
    - 1242×2208 竖版 (iPhone 6/7/8 Plus 逻辑分辨率, 移动端阅读舒适)
    - 大标题、舒适行距、深灰底色护眼
    - 每页智能切分, 不截断句子
    - 页底页码 + 进度提示
"""

import os
import sys
import re
import argparse
from xml.etree import ElementTree as ET

# ── 依赖检查 ──────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, Emu
except ImportError:
    sys.exit("❌ 缺少 python-docx: pip3 install python-docx")

try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    sys.exit("❌ 缺少 Pillow: pip3 install Pillow")
# FreeTypeFont type hint (some Pillow versions expose Font but not FreeTypeFont)
FontType = ImageFont.FreeTypeFont if hasattr(ImageFont, 'FreeTypeFont') else ImageFont.ImageFont


# ── 默认配置 ──────────────────────────────────────────────────
CANVAS_W = 1242
CANVAS_H = 2208

MARGIN_LR = 80   # 左右边距
MARGIN_TB = 120  # 上下边距

TEXT_W = CANVAS_W - MARGIN_LR * 2
TEXT_H = CANVAS_H - MARGIN_TB * 2

# 行高倍数 (相对于字号)
LINE_H = 1.7

# 字体大小 (不同元素)
SIZE_H1 = 54
SIZE_H2 = 44
SIZE_H3 = 38
SIZE_BODY = 34
SIZE_LIST = 34
SIZE_PAGE = 26   # 页码

# 颜色
BG_COLOR = (247, 243, 236)     # 暖白仿纸
TEXT_COLOR = (45, 40, 35)      # 深灰
ACCENT_COLOR = (180, 80, 50)   # 强调色 (标题装饰线)
PAGE_NUM_COLOR = (170, 165, 155)

# 标题前的装饰条宽
ACCENT_BAR_W = 6


def resolve_font(size: int, style: str = "regular") -> ImageFont.FreeTypeFont:
    """智能选择中文字体, 按优先级降序."""
    candidates = {
        "regular": [
            "/System/Library/Fonts/PingFang.ttc",           # macOS
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/System/Library/Fonts/STHeiti Medium.ttc",
            "/System/Library/Fonts/Hiragino Sans GB.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",  # Linux
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        ],
        "bold": [
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Medium.ttc",
            "/System/Library/Fonts/Hiragino Sans GB W6.ttc",
            "/System/Library/Fonts/Hiragino Sans GB.ttc",
        ],
    }
    for path in candidates.get(style, candidates["regular"]):
        if os.path.exists(path):
            try:
                return ImageFont.truetype(path, size)
            except Exception:
                continue
    # Fallback: PIL 默认
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.Draw, text: str, font, max_w: int) -> list[str]:
    """按宽度自动换行, 返回行列表."""
    lines = []
    for paragraph in text.split("\n"):
        if not paragraph:
            lines.append("")
            continue
        words = paragraph
        line = ""
        for ch in words:
            test = line + ch
            bbox = draw.textbbox((0, 0), test, font=font)
            if bbox[2] - bbox[0] <= max_w:
                line = test
            else:
                lines.append(line)
                line = ch
        if line:
            lines.append(line)
    return lines


class WordRenderer:
    """读取 docx → 分成页面 → 渲染为 PIL Image 列表"""

    def __init__(self, docx_path: str):
        self.doc = Document(docx_path)
        self.font_body = resolve_font(SIZE_BODY)
        self.font_h1 = resolve_font(SIZE_H1, "bold")
        self.font_h2 = resolve_font(SIZE_H2, "bold")
        self.font_h3 = resolve_font(SIZE_H3, "bold")
        self.font_list = resolve_font(SIZE_LIST)
        self.font_page = resolve_font(SIZE_PAGE)

    def extract_blocks(self) -> list[dict]:
        """提取文档中的文本块, 每块带类型."""
        blocks = []
        numbering = self._get_ordered_list_numbers()

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name.lower() if para.style else ""

            block = {"text": text, "type": "body", "bold": False}

            # 判断类型
            if style_name.startswith("heading 1") or style_name == "标题 1":
                block["type"] = "h1"
            elif style_name.startswith("heading 2") or style_name == "标题 2":
                block["type"] = "h2"
            elif style_name.startswith("heading 3") or style_name == "标题 3":
                block["type"] = "h3"
            elif style_name.startswith("list") or style_name == "列表段落":
                # 检查是否有编号
                num_id = para._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')
                if num_id is not None and i in numbering:
                    block["type"] = "ordered_list"
                    block["prefix"] = f"{numbering[i]}. "
                else:
                    block["type"] = "list"
                    block["prefix"] = "• "
            else:
                # 检查是否加粗
                for run in para.runs:
                    if run.bold:
                        block["bold"] = True
                        break

            blocks.append(block)

        return blocks

    def _get_ordered_list_numbers(self) -> dict:
        """尝试解析 docx 中的有序列表编号 (近似)."""
        # 简单实现: 连续 list 段落累加编号
        numbers = {}
        num = 0
        in_list = False
        for i, para in enumerate(self.doc.paragraphs):
            style_name = para.style.name.lower() if para.style else ""
            if style_name.startswith("list"):
                # 检查 XML 是否有 numPr
                ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
                num_pr = para._element.find(f'.//{ns}numPr')
                if num_pr is not None:
                    in_list = True
                    num += 1
                    numbers[i] = num
                else:
                    # 无序列表
                    pass
            else:
                in_list = False
                num = 0
        return numbers

    def page_blocks(self, blocks: list[dict]) -> list[list[dict]]:
        """将文本块分页, 保证不截断句子."""
        pages = []
        current_page = []
        current_h = MARGIN_TB  # 已使用的垂直空间

        def line_height(size):
            return int(size * LINE_H)

        def block_height(block):
            btype = block["type"]
            if btype == "h1":
                return line_height(SIZE_H1) + 20
            elif btype == "h2":
                return line_height(SIZE_H2) + 16
            elif btype == "h3":
                return line_height(SIZE_H3) + 12
            else:
                return line_height(SIZE_BODY) + 6

        # 先展开每块为预渲染行
        temp_font_body = self.font_body
        draw = ImageDraw.Draw(Image.new("RGB", (CANVAS_W, CANVAS_H)))
        for block in blocks:
            btype = block["type"]
            if btype == "h1":
                font = self.font_h1
            elif btype == "h2":
                font = self.font_h2
            elif btype == "h3":
                font = self.font_h3
            else:
                font = self.font_body

            prefix = block.get("prefix", "")
            lines = wrap_text(draw, prefix + block["text"], font, TEXT_W)
            block["_lines"] = lines
            block["_font"] = font

            if btype == "h1":
                block["_lh"] = line_height(SIZE_H1)
            elif btype == "h2":
                block["_lh"] = line_height(SIZE_H2)
            elif btype == "h3":
                block["_lh"] = line_height(SIZE_H3)
            else:
                block["_lh"] = line_height(SIZE_BODY)

        # 分页
        for block in blocks:
            bh = block["_lh"] * len(block["_lines"]) + 20  # +段间距
            if current_h + bh > CANVAS_H - MARGIN_TB:
                if current_page:
                    pages.append(current_page)
                current_page = [block]
                current_h = MARGIN_TB + bh
            else:
                current_page.append(block)
                current_h += bh

        if current_page:
            pages.append(current_page)

        return pages

    def render_page(self, blocks: list[dict], page_num: int, total: int) -> Image.Image:
        """渲染单页为 PIL Image."""
        img = Image.new("RGB", (CANVAS_W, CANVAS_H), BG_COLOR)
        draw = ImageDraw.Draw(img)

        y = MARGIN_TB

        for block in blocks:
            btype = block["type"]
            font = block["_font"]
            lines = block["_lines"]
            lh = block["_lh"]
            prefix = block.get("prefix", "")

            if btype == "h1":
                # 左侧装饰条 + 标题
                draw.rectangle(
                    [MARGIN_LR, y + 8, MARGIN_LR + ACCENT_BAR_W, y + lh * len(lines) - 8],
                    fill=ACCENT_COLOR,
                )
                x_text = MARGIN_LR + ACCENT_BAR_W + 20
                for line in lines:
                    draw.text((x_text, y), line, font=font, fill=TEXT_COLOR)
                    y += lh
                y += 20

            elif btype == "h2":
                for line in lines:
                    draw.text((MARGIN_LR + 10, y), line, font=font, fill=TEXT_COLOR)
                    y += lh
                # 下划线装饰
                draw.line(
                    [MARGIN_LR, y - 6, MARGIN_LR + 200, y - 6],
                    fill=ACCENT_COLOR, width=3,
                )
                y += 16

            elif btype in ("list", "ordered_list"):
                for line in lines:
                    draw.text((MARGIN_LR + 10, y), line, font=font, fill=TEXT_COLOR)
                    y += lh
                y += 6

            else:
                for line in lines:
                    draw.text((MARGIN_LR + 10, y), line, font=font, fill=TEXT_COLOR)
                    y += lh
                y += 8

        # ── 页码 ──
        page_text = f"— {page_num}/{total} —"
        pb = draw.textbbox((0, 0), page_text, font=self.font_page)
        pw = pb[2] - pb[0]
        draw.text(
            ((CANVAS_W - pw) // 2, CANVAS_H - MARGIN_TB // 2),
            page_text, font=self.font_page, fill=PAGE_NUM_COLOR,
        )

        return img


def main():
    parser = argparse.ArgumentParser(description="Word 文档 → 多张手机端图片")
    parser.add_argument("input", help="输入 .docx 文件路径")
    parser.add_argument("--output-dir", "-o",
                        default=os.path.join(os.path.dirname(os.path.abspath(__file__)), "output_images"),
                        help="输出目录 (默认: 脚本所在目录/output_images)")
    parser.add_argument("--format", "-f", choices=["png", "jpg"], default="png",
                        help="图片格式 (默认: png)")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        sys.exit(f"❌ 文件不存在: {args.input}")
    if not args.input.endswith(".docx"):
        sys.exit("❌ 仅支持 .docx 格式")

    os.makedirs(args.output_dir, exist_ok=True)

    base = os.path.splitext(os.path.basename(args.input))[0]
    fmt = args.format

    print(f"📄 读取: {args.input}")
    renderer = WordRenderer(args.input)

    print("🔍 提取文本块...")
    blocks = renderer.extract_blocks()
    print(f"   共 {len(blocks)} 个文本块")

    print("📄 分页...")
    pages = renderer.page_blocks(blocks)
    total = len(pages)
    print(f"   共 {total} 页")

    for i, page_blocks in enumerate(pages, 1):
        out_path = os.path.join(args.output_dir, f"{base}_p{i:02d}.{fmt}")
        img = renderer.render_page(page_blocks, i, total)
        img.save(out_path, quality=95 if fmt == "jpg" else None)
        print(f"   ✅ 第 {i}/{total} 页 → {out_path}")

    print(f"\n🎉 完成! 共生成 {total} 张图片到 {args.output_dir}/")


if __name__ == "__main__":
    main()
